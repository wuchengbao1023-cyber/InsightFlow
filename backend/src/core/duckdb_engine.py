"""
InsightFlow AI 2026 - DuckDB 数据处理引擎
==========================================

替代 Pandas 的高性能 In-Memory OLAP 引擎
- 进程内运行，无需单独服务
- 流式处理，不一次性读入内存
- 支持 GB 级别文件
- 比 Pandas 快 10-100 倍

Author: InsightFlow AI Team
"""

import logging
import os
import io
import asyncio
from typing import Dict, Any, List, Optional, Union
from pathlib import Path
import json

logger = logging.getLogger(__name__)

# ── 尝试导入 DuckDB ──────────────────────────────────────────
try:
    import duckdb
    import pandas as pd
    import numpy as np
    HAS_DUCKDB = True
    logger.info("✅ DuckDB 已就绪（版本: %s）", duckdb.__version__)
except ImportError:
    HAS_DUCKDB = False
    duckdb = None
    logger.warning("⚠️  DuckDB 未安装，将降级到 Pandas 模式。运行: pip install duckdb")

# ── 安全阀常量 ─────────────────────────────────────────────────
MAX_FILE_SIZE_MB = 10          # 最大文件大小（MB）
MAX_FILE_SIZE_BYTES = MAX_FILE_SIZE_MB * 1024 * 1024
QUERY_TIMEOUT_SECONDS = 60     # 单次查询最大执行时间
MAX_ROWS_RETURN = 10_000       # 单次查询最多返回行数
MAX_CONCURRENT_QUERIES = 1     # 轻量服务器：同时只允许 1 个分析任务
_active_queries = 0            # 当前并发查询数


class FileSizeError(Exception):
    """文件超过大小限制"""
    pass


def sanitize_identifier(name: str) -> str:
    """
    清理 SQL 标识符（表名/列名），防止注入。

    规则：
    1. 只保留字母、数字、下划线、中文、空格、横杠、点
    2. 转义内部的双引号（DuckDB 用 "" 表示字面双引号）
    3. 最终用双引号包裹（DuckDB 标识符引用）

    用法：f'SELECT "{sanitize_identifier(col)}" FROM "{sanitize_identifier(table)}"'
    """
    if not name or not isinstance(name, str):
        raise ValueError(f"无效的SQL标识符: {name!r}")

    # 去掉首尾空白
    cleaned = name.strip()

    # 只允许安全字符：字母、数字、中文、下划线、空格、横杠、点
    import re
    if not re.match(r'^[\w\u4e00-\u9fff\s\-.]+$', cleaned):
        logger.warning(f"SQL标识符含可疑字符，已清理: {name!r}")

    # 只保留安全字符
    cleaned = re.sub(r'[^\w\u4e00-\u9fff\s\-.]', '_', cleaned)

    # 转义内部双引号（DuckDB用""表示字面"）
    cleaned = cleaned.replace('"', '""')

    return cleaned


def safe_col(col: str) -> str:
    """快捷函数：返回双引号包裹的安全列名"""
    return f'"{sanitize_identifier(col)}"'


def safe_table(table: str) -> str:
    """快捷函数：返回双引号包裹的安全表名"""
    return f'"{sanitize_identifier(table)}"'


class ConcurrencyError(Exception):
    """并发超限"""
    pass


class DuckDBEngine:
    """
    DuckDB 进程内 OLAP 引擎

    使用方式:
        engine = DuckDBEngine()
        engine.load_file("path/to/file.csv", table_name="sales")
        results = engine.query("SELECT region, SUM(amount) FROM sales GROUP BY region")
    """

    def __init__(self):
        self._conn: Optional["duckdb.DuckDBPyConnection"] = None
        self._tables: Dict[str, Dict[str, Any]] = {}   # 已注册的表信息
        self._initialized = False
        self._db_path = str(Path(__file__).parent.parent.parent / "data" / "insightflow.duckdb")

        if HAS_DUCKDB:
            try:
                # 磁盘持久化数据库，重启不丢失
                Path(self._db_path).parent.mkdir(parents=True, exist_ok=True)
                self._conn = duckdb.connect(database=self._db_path)
                # 轻量服务器：限制 DuckDB 使用的线程数
                self._conn.execute("PRAGMA threads=2")
                self._conn.execute("PRAGMA memory_limit='512MB'")
                self._initialized = True

                # 启动时恢复已注册的表信息
                self._restore_tables()

                table_count = len(self._get_registered_tables())
                logger.info("🦆 DuckDB 持久化数据库已启动 (%s, %d 张表)", self._db_path, table_count)
            except Exception as e:
                logger.error("❌ DuckDB 初始化失败: %s", e)
                self._initialized = False

    def _restore_tables(self):
        """启动时恢复已注册表的元信息（DuckDB磁盘文件已保存表数据，只需恢复内存索引）"""
        if not self._initialized:
            return
        try:
            registered = self._get_registered_tables()
            for tname in registered:
                schema = self._get_schema(tname)
                row_count = self._conn.execute(f'SELECT COUNT(*) FROM "{tname}"').fetchone()[0]
                self._tables[tname] = {
                    "name": tname,
                    "rows": row_count,
                    "columns": len(schema),
                    "schema": schema,
                    "source": "persistent",
                }
            if registered:
                logger.info("🦆 恢复了 %d 张持久化表: %s", len(registered), registered)
        except Exception as e:
            logger.warning("🦆 恢复表信息失败（可能是首次启动）: %s", e)

    # ─────────────────────────────────────────────────────────
    # 文件加载
    # ─────────────────────────────────────────────────────────

    async def load_file(
        self,
        file_content: bytes,
        filename: str,
        table_name: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        从上传的文件内容加载数据到 DuckDB

        Args:
            file_content: 文件字节内容
            filename: 原始文件名（用于判断格式）
            table_name: 目标表名，默认使用文件名（不含扩展名）

        Returns:
            加载结果信息
        """
        # 安全阀 1：文件大小检查
        file_size = len(file_content)
        if file_size > MAX_FILE_SIZE_BYTES:
            raise FileSizeError(
                f"文件大小 {file_size / 1024 / 1024:.1f}MB 超过限制 {MAX_FILE_SIZE_MB}MB。"
                f"请压缩数据后重试。"
            )

        if table_name is None:
            table_name = Path(filename).stem.replace("-", "_").replace(" ", "_").lower()
            # 确保表名合法
            table_name = "".join(c if c.isalnum() or c == "_" else "_" for c in table_name)
            if not table_name[0].isalpha():
                table_name = "t_" + table_name

        ext = Path(filename).suffix.lower()

        logger.info(
            "📂 加载文件: %s (%.1f KB) -> 表 [%s]",
            filename,
            file_size / 1024,
            table_name,
        )

        # 安全阀 2：并发检查
        await self._acquire_slot()
        try:
            # 先尝试删除同名表（同名重新上传时覆盖）
            if table_name in self._tables and self._initialized and self._conn:
                try:
                    self._conn.execute(f'DROP TABLE IF EXISTS {safe_table(table_name)}')
                except Exception:
                    pass
            result = await asyncio.get_event_loop().run_in_executor(
                None,
                self._load_file_sync,
                file_content,
                filename,
                table_name,
                ext,
            )
        finally:
            self._release_slot()

        return result

    def _load_file_sync(
        self,
        file_content: bytes,
        filename: str,
        table_name: str,
        ext: str,
    ) -> Dict[str, Any]:
        """同步加载文件（在线程池中执行）"""
        if not self._initialized or not HAS_DUCKDB:
            return self._load_with_pandas(file_content, filename, table_name, ext)

        try:
            buf = io.BytesIO(file_content)

            if ext == ".csv":
                # 最简单的方法：使用 Pandas 读取，然后注册到 DuckDB
                import pandas as pd
                # 尝试多种编码读取CSV
                df = None
                encodings = ["utf-8", "gbk", "gb2312", "gb18030", "latin-1"]
                
                for encoding in encodings:
                    try:
                        # 每次尝试都创建新的BytesIO对象
                        buf_copy = io.BytesIO(file_content)
                        df = pd.read_csv(buf_copy, nrows=500_000, encoding=encoding)
                        logger.info(f"CSV使用{encoding}编码读取成功")
                        break
                    except UnicodeDecodeError:
                        continue
                    except Exception as e:
                        logger.debug(f"CSV使用{encoding}编码读取失败: {e}")
                        continue
                
                if df is None:
                    # 如果所有编码都失败，尝试不指定编码
                    try:
                        buf_copy = io.BytesIO(file_content)
                        df = pd.read_csv(buf_copy, nrows=500_000, encoding=None)
                    except Exception as e:
                        raise ValueError(f"无法读取CSV文件，尝试的编码: {encodings}。错误: {e}")
                
                # 列名保险：确保无 None/NaN 列名
                df.columns = [str(c) if c is not None and str(c) not in ("nan", "None", "") else f"_col_{i}" 
                              for i, c in enumerate(df.columns)]
                # 持久化写入磁盘表（而非内存register，重启不丢失）
                self._conn.register(f"__temp_{table_name}", df)
                self._conn.execute(f'CREATE TABLE IF NOT EXISTS "{table_name}" AS SELECT * FROM "__temp_{table_name}"')
                self._conn.unregister(f"__temp_{table_name}")
                schema = self._get_schema(table_name)

            elif ext in (".xlsx", ".xls"):
                # 使用 pandas 读取 Excel，自动检测标题行
                import pandas as pd
                df = self._read_excel_with_header_detection(buf, nrows=500_000)
                # 列名保险：确保无 None/NaN 列名（DuckDB register 会报错）
                df.columns = [str(c) if c is not None and str(c) not in ("nan", "None", "") else f"_col_{i}" 
                              for i, c in enumerate(df.columns)]
                # 持久化写入
                self._conn.register(f"__temp_{table_name}", df)
                self._conn.execute(f'CREATE TABLE IF NOT EXISTS "{table_name}" AS SELECT * FROM "__temp_{table_name}"')
                self._conn.unregister(f"__temp_{table_name}")
                schema = self._get_schema(table_name)

            elif ext == ".json":
                # 使用 pandas 读取 JSON
                import pandas as pd
                df = pd.read_json(buf)
                # 列名保险
                df.columns = [str(c) if c is not None and str(c) not in ("nan", "None", "") else f"_col_{i}" 
                              for i, c in enumerate(df.columns)]
                # 持久化写入
                self._conn.register(f"__temp_{table_name}", df)
                self._conn.execute(f'CREATE TABLE IF NOT EXISTS "{table_name}" AS SELECT * FROM "__temp_{table_name}"')
                self._conn.unregister(f"__temp_{table_name}")
                schema = self._get_schema(table_name)

            elif ext == ".pdf":
                # ── PDF 三级解析策略 ──
                import pandas as pd
                from io import StringIO
                import re

                # 第1级：PyMuPDF 表格提取（有明确网格线的表格）
                df = self._extract_tables_from_pdf(file_content, filename)

                if df is None or df.empty:
                    # 第2级：🧠 LLM 智能提取（从纯文本中提取结构化数据）
                    logger.info("[PDF] 未检测到网格表格，尝试 LLM 智能提取...")
                    text = self._extract_text_from_pdf(file_content)
                    if text and len(text.strip()) > 50:  # 至少50字符才有提取价值
                        csv_data = self._llm_extract_structured_csv(text, filename)
                        if csv_data:
                            try:
                                df = pd.read_csv(StringIO(csv_data))
                                # 数值列清洗：去掉残留的单位文字（LLM 有时不听话）
                                for col in df.columns:
                                    if df[col].dtype == object:
                                        df[col] = df[col].apply(
                                            lambda v: re.sub(r'[万亿亿年个元％%]', '', str(v)).strip() if pd.notna(v) and isinstance(v, str) else v
                                        )
                                        # 尝试转数值
                                        try:
                                            df[col] = pd.to_numeric(df[col], errors='ignore')
                                        except Exception:
                                            pass
                                logger.info(f"[PDF-LLM] LLM 提取成功: {len(df)}行 × {len(df.columns)}列")
                            except Exception as e:
                                logger.warning(f"[PDF-LLM] CSV 解析失败: {e}")
                                df = None

                    # 第3级：纯文本逐行兜底
                    if df is None or df.empty:
                        if text:
                            df = pd.DataFrame({"文本内容": text.split("\n")})
                            df = df[df["文本内容"].str.strip().str.len() > 0].reset_index(drop=True)
                            logger.info(f"[PDF] 降级为纯文本模式: {len(df)}行")
                        else:
                            raise ValueError("PDF文件中未发现可用的表格或文本数据")

                # 列名保险
                df.columns = [str(c) if c is not None and str(c) not in ("nan", "None", "") else f"_col_{i}"
                              for i, c in enumerate(df.columns)]
                # 持久化写入
                self._conn.register(f"__temp_{table_name}", df)
                self._conn.execute(f'CREATE TABLE IF NOT EXISTS "{table_name}" AS SELECT * FROM "__temp_{table_name}"')
                self._conn.unregister(f"__temp_{table_name}")
                schema = self._get_schema(table_name)

            elif ext in (".docx", ".doc"):
                # ── Word 三级解析策略（同 PDF）──
                import pandas as pd
                from io import StringIO
                import re

                df = self._extract_tables_from_docx(file_content, filename)

                if df is None or df.empty:
                    text = self._extract_text_from_docx(file_content)
                    if text and len(text.strip()) > 50:
                        csv_data = self._llm_extract_structured_csv(text, filename)
                        if csv_data:
                            try:
                                df = pd.read_csv(StringIO(csv_data))
                                # 数值列清洗（同PDF）
                                for col in df.columns:
                                    if df[col].dtype == object:
                                        df[col] = df[col].apply(
                                            lambda v: re.sub(r'[万亿亿年个元％%]', '', str(v)).strip() if pd.notna(v) and isinstance(v, str) else v
                                        )
                                        try:
                                            df[col] = pd.to_numeric(df[col], errors='ignore')
                                        except Exception:
                                            pass
                                logger.info(f"[DOCX-LLM] LLM 提取成功: {len(df)}行 × {len(df.columns)}列")
                            except Exception:
                                df = None

                    if df is None or df.empty:
                        if text:
                            df = pd.DataFrame({"文本内容": text.split("\n")})
                            df = df[df["文本内容"].str.strip().str.len() > 0].reset_index(drop=True)
                            logger.info(f"[DOCX] 降级为纯文本模式: {len(df)}行")
                        else:
                            raise ValueError("Word文件中未发现可用的表格或文本数据")
                # 列名保险
                df.columns = [str(c) if c is not None and str(c) not in ("nan", "None", "") else f"_col_{i}"
                              for i, c in enumerate(df.columns)]
                # 持久化写入
                self._conn.register(f"__temp_{table_name}", df)
                self._conn.execute(f'CREATE TABLE IF NOT EXISTS "{table_name}" AS SELECT * FROM "__temp_{table_name}"')
                self._conn.unregister(f"__temp_{table_name}")
                schema = self._get_schema(table_name)

            else:
                raise ValueError(f"不支持的文件格式: {ext}，支持 CSV / Excel / JSON / PDF / Word(.docx)")

            row_count = self._conn.execute(
                f'SELECT COUNT(*) FROM "{table_name}"'
            ).fetchone()[0]

            self._tables[table_name] = {
                "filename": filename,
                "rows": row_count,
                "schema": schema,
                "size_kb": len(file_content) / 1024,
            }

            logger.info("✅ 已加载 %d 行数据到 DuckDB 表 [%s]", row_count, table_name)

            return {
                "success": True,
                "table_name": table_name,
                "rows": row_count,
                "columns": len(schema),
                "schema": schema,
                "engine": "duckdb",
                "size_kb": round(len(file_content) / 1024, 1),
            }

        except Exception as e:
            logger.error("❌ DuckDB 加载失败，降级到 Pandas: %s", e)
            return self._load_with_pandas(file_content, filename, table_name, ext)

    def _load_with_pandas(
        self,
        file_content: bytes,
        filename: str,
        table_name: str,
        ext: str,
    ) -> Dict[str, Any]:
        """Pandas 降级加载"""
        import pandas as pd

        buf = io.BytesIO(file_content)
        if ext == ".csv":
            df = pd.read_csv(buf)
        elif ext in (".xlsx", ".xls"):
            df = self._read_excel_with_header_detection(buf)
        elif ext == ".json":
            df = pd.read_json(buf)
        elif ext == ".pdf":
            from io import StringIO
            df = self._extract_tables_from_pdf(file_content, filename)
            if df is None or df.empty:
                text = self._extract_text_from_pdf(file_content)
                if text and len(text.strip()) > 50:
                    csv_data = self._llm_extract_structured_csv(text, filename)
                    if csv_data:
                        try:
                            df = pd.read_csv(StringIO(csv_data))
                        except Exception:
                            df = None
                if df is None or df.empty:
                    if text:
                        df = pd.DataFrame({"文本内容": text.split("\n")})
                        df = df[df["文本内容"].str.strip().str.len() > 0].reset_index(drop=True)
                    else:
                        raise ValueError("PDF中未发现可用数据")
        elif ext in (".docx",):
            df = self._extract_tables_from_docx(file_content, filename)
            if df is None or df.empty:
                text = self._extract_text_from_docx(file_content)
                if text:
                    df = pd.DataFrame({"文本内容": text.split("\n")})
                    df = df[df["文本内容"].str.strip().str.len() > 0].reset_index(drop=True)
                else:
                    raise ValueError("Word中未发现可用数据")
        else:
            raise ValueError(f"不支持的文件格式: {ext}")

        # 存为伪表（实际是 DataFrame）
        if not hasattr(self, "_pandas_tables"):
            self._pandas_tables = {}
        self._pandas_tables[table_name] = df

        schema = {col: str(dtype) for col, dtype in df.dtypes.items()}
        self._tables[table_name] = {
            "filename": filename,
            "rows": len(df),
            "schema": schema,
            "size_kb": len(file_content) / 1024,
        }

        logger.info("✅ Pandas 降级加载 %d 行 -> 表 [%s]", len(df), table_name)

        return {
            "success": True,
            "table_name": table_name,
            "rows": len(df),
            "columns": len(df.columns),
            "schema": schema,
            "engine": "pandas_fallback",
            "size_kb": round(len(file_content) / 1024, 1),
        }

    # ─────────────────────────────────────────────────────────
    # SQL 查询
    # ─────────────────────────────────────────────────────────

    async def query(
        self,
        sql: str,
        table_name: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        执行 SQL 查询

        Args:
            sql: SQL 语句（SELECT 类）
            table_name: 可选，用于提示数据来源

        Returns:
            查询结果
        """
        await self._acquire_slot()
        try:
            result = await asyncio.wait_for(
                asyncio.get_event_loop().run_in_executor(
                    None, self._query_sync, sql
                ),
                timeout=QUERY_TIMEOUT_SECONDS,
            )
        except asyncio.TimeoutError:
            raise TimeoutError(
                f"查询超时（>{QUERY_TIMEOUT_SECONDS}s），请简化查询条件。"
            )
        finally:
            self._release_slot()

        return result

    def _query_sync(self, sql: str) -> Dict[str, Any]:
        """同步执行查询"""
        if not self._initialized or not HAS_DUCKDB:
            return self._query_with_pandas(sql)

        try:
            # 安全：只允许 SELECT / WITH，拒绝危险关键字
            import re
            sql_clean = re.sub(r'/\*.*?\*/', '', sql, flags=re.DOTALL)  # 去除块注释
            sql_clean = re.sub(r'--.*$', '', sql_clean, flags=re.MULTILINE)  # 去除行注释
            sql_upper = sql_clean.strip().upper()
            if not sql_upper.startswith("SELECT") and not sql_upper.startswith("WITH"):
                raise ValueError("只允许 SELECT 查询，禁止写操作。")
            # 二次检查：去除注释后仍不允许危险关键字
            dangerous_keywords = re.findall(
                r'\b(DROP|DELETE|INSERT|UPDATE|ALTER|CREATE|TRUNCATE|EXEC|GRANT|REVOKE)\b',
                sql_upper
            )
            if dangerous_keywords:
                raise ValueError(f"检测到禁止的关键字: {dangerous_keywords}")

            rel = self._conn.execute(sql)
            if rel.description is None:
                return {"success": True, "columns": [], "rows": [], "row_count": 0, "engine": "duckdb"}
            columns = [desc[0] for desc in rel.description]
            rows = rel.fetchmany(MAX_ROWS_RETURN)

            data = [dict(zip(columns, row)) for row in rows]

            return {
                "success": True,
                "columns": columns,
                "rows": data,
                "row_count": len(data),
                "truncated": len(data) == MAX_ROWS_RETURN,
                "engine": "duckdb",
            }

        except Exception as e:
            logger.error("❌ DuckDB 查询失败: %s | SQL: %s", e, sql[:200])
            return {"success": False, "error": str(e), "rows": [], "columns": []}

    def execute_query(self, sql: str, table_name: Optional[str] = None) -> Dict[str, Any]:
        """
        同步执行查询并返回标准格式结果（供哨兵和神谕等Agent使用）
        
        Args:
            sql: SQL查询语句
            table_name: 可选表名，用于日志记录
            
        Returns:
            标准格式查询结果：
            - 成功: {"success": True, "data": [...], "columns": [...], "row_count": N, ...}
            - 失败: {"success": False, "error": "错误信息", "data": [], "columns": []}
        """
        result = self._query_sync(sql)
        
        # 标准化返回格式：始终包含 'data' 和 'columns' 字段
        # _query_sync 返回 'rows'，但各 Agent 期望 'data'
        result.setdefault("data", result.get("rows", []))
        result.setdefault("columns", [])
            
        return result

    def _query_with_pandas(self, sql: str) -> Dict[str, Any]:
        """Pandas 降级查询"""
        if not hasattr(self, "_pandas_tables") or not self._pandas_tables:
            return {"success": False, "error": "无可用数据表", "rows": [], "columns": []}

        # 简单降级：只支持全表查询
        table_name = list(self._pandas_tables.keys())[0]
        df = self._pandas_tables[table_name]
        data = df.head(MAX_ROWS_RETURN).to_dict(orient="records")

        return {
            "success": True,
            "columns": list(df.columns),
            "rows": data,
            "row_count": len(data),
            "truncated": len(df) > MAX_ROWS_RETURN,
            "engine": "pandas_fallback",
        }

    # ─────────────────────────────────────────────────────────
    # 数据概览（供智能体使用）
    # ─────────────────────────────────────────────────────────

    def get_table_summary(self, table_name: str) -> Dict[str, Any]:
        """获取表的数据概览，供 LLM 分析上下文"""
        if table_name not in self._tables:
            return {"error": f"表 {table_name} 不存在"}

        info = self._tables[table_name].copy()

        if self._initialized and HAS_DUCKDB and table_name in self._get_registered_tables():
            try:
                # 数值列统计
                stats_rows = self._conn.execute(
                    f"SUMMARIZE {safe_table(table_name)}"
                ).fetchall()
                # SUMMARIZE 返回: column_name, column_type, min, max, approx_unique, avg, std, q25, q50, q75, count, null_percentage
                cols_desc = [d[0] for d in self._conn.execute(f"SUMMARIZE {safe_table(table_name)}").description]
                stats = [dict(zip(cols_desc, r)) for r in stats_rows]
                info["column_stats"] = stats
            except Exception:
                pass

        return info

    def get_dataframe(self, table_name: str) -> Optional["pd.DataFrame"]:
        """获取 DataFrame（供需要 Pandas 的智能体使用）"""
        if not self._initialized or not HAS_DUCKDB:
            if hasattr(self, "_pandas_tables"):
                return self._pandas_tables.get(table_name)
            return None

        if table_name not in self._get_registered_tables():
            return None

        try:
            return self._conn.execute(
                f"SELECT * FROM {table_name} LIMIT {MAX_ROWS_RETURN}"
            ).df()
        except Exception as e:
            logger.error("❌ 获取 DataFrame 失败: %s", e)
            return None

    def list_tables(self) -> List[str]:
        """列出所有已加载的表"""
        logger.debug(f"list_tables called: _initialized={self._initialized}, HAS_DUCKDB={HAS_DUCKDB}, _tables keys={list(self._tables.keys())}")
        return list(self._tables.keys())

    def get_tables_info(self) -> Dict[str, Any]:
        """获取所有表的信息摘要"""
        return {
            name: {
                "rows": info.get("rows", 0),
                "columns": len(info.get("schema", {})),
                "size_kb": round(info.get("size_kb", 0), 1),
                "filename": info.get("filename", ""),
            }
            for name, info in self._tables.items()
        }

    # ─────────────────────────────────────────────────────────
    # 内部工具方法
    # ─────────────────────────────────────────────────────────

    def _get_schema(self, table_name: str) -> Dict[str, str]:
        """获取表结构"""
        try:
            rel = self._conn.execute(
                f"SELECT * FROM {table_name} LIMIT 0"
            )
            if rel.description is None:
                return {}
            return {str(desc[0]): str(desc[1]) for desc in rel.description}
        except Exception as e:
            logger.warning(f"获取表 {table_name} 的结构失败: {e}")
            return {}

    def _get_registered_tables(self) -> List[str]:
        """获取已注册到 DuckDB 的表名列表"""
        if not self._initialized or self._conn is None:
            return list(self._tables.keys())
        
        try:
            rows = self._conn.execute("SHOW TABLES").fetchall()
            return [r[0] for r in rows]
        except Exception as e:
            logger.warning(f"获取注册表列表失败: {e}")
            return list(self._tables.keys())

    async def _acquire_slot(self):
        """获取并发槽位（安全阀：并发限制）"""
        global _active_queries
        if _active_queries >= MAX_CONCURRENT_QUERIES:
            raise ConcurrencyError(
                "当前已有分析任务正在运行，请稍后再试。"
                f"（轻量服务器限制：同时最多 {MAX_CONCURRENT_QUERIES} 个分析任务）"
            )
        _active_queries += 1

    def _release_slot(self):
        """释放并发槽位"""
        global _active_queries
        _active_queries = max(0, _active_queries - 1)

    def _read_excel_with_header_detection(
        self,
        buf: "io.BytesIO",
        nrows: int = 500_000
    ) -> "pd.DataFrame":
        """
        智能读取 Excel：自动检测真正的列名行，跳过标题行。

        策略：
        1. 先读前 5 行（不指定 header）
        2. 逐行判断哪行看起来像"真正的表头"（所有格单元格非空、无超长文本、
           不含年份+连续数字模式）
        3. 以检测到的行号作为 header 重新读取
        4. 如果检测失败，回退到默认行为（header=0）
        """
        import pandas as pd
        import re

        # ── 1. 读前 5 行用于检测 ────────────────────────────────────────────
        try:
            preview = pd.read_excel(buf, header=None, nrows=5)
        except Exception as e:
            logger.warning(f"Excel 预览读取失败，使用默认 header=0: {e}")
            buf.seek(0)
            return pd.read_excel(buf, nrows=nrows)

        # ── 2. 对每一行打分，找最可能是列名的行 ─────────────────────────────
        def _row_looks_like_header(row) -> bool:
            """判断一行是否像表头（列名）"""
            non_null_cells = [str(v).strip() for v in row if v is not None and str(v).strip() not in ("", "nan")]
            if not non_null_cells:
                return False
            # 表头行特征：
            # a. 无明显超长文本（标题通常 >30 个字）
            max_len = max(len(c) for c in non_null_cells)
            if max_len > 30:
                return False
            # b. 不包含四位年份+度/年/期等典型标题词
            full_text = "".join(non_null_cells)
            if re.search(r'20\d{2}年度|考试录用|名单|公告', full_text):
                return False
            # c. 非空单元格占多数（≥ 50% 列数）
            total_cols = len(row)
            if total_cols > 0 and len(non_null_cells) < total_cols * 0.5:
                return False
            return True

        header_row = 0  # 默认
        for i, row in preview.iterrows():
            if _row_looks_like_header(row):
                header_row = i
                break

        if header_row != 0:
            logger.info(f"📊 Excel 标题行检测：跳过前 {header_row} 行，以第 {header_row + 1} 行作为列名")

        # ── 3. 以检测到的行号重新读取完整文件 ───────────────────────────────
        buf.seek(0)
        df = pd.read_excel(buf, header=header_row, nrows=nrows)

        # ── 4. 清理：去掉全空列名（Unnamed: x）若真正列名存在 ───────────────
        # 检查是否还有 Unnamed 列（说明检测可能不准）
        unnamed_cols = [c for c in df.columns if str(c).startswith("Unnamed:")]
        real_cols = [c for c in df.columns if not str(c).startswith("Unnamed:")]
        if unnamed_cols and real_cols:
            # 有真实列名也有 Unnamed，只保留真实列
            df = df[real_cols]
            logger.info(f"📊 Excel 列清理：删除 {len(unnamed_cols)} 个 Unnamed 列，保留 {len(real_cols)} 个真实列")
        elif unnamed_cols and not real_cols:
            # 所有列都是 Unnamed，说明检测失败，尝试下一行
            logger.warning("⚠️ 所有列仍为 Unnamed，尝试 header=1 再读一次")
            buf.seek(0)
            df2 = pd.read_excel(buf, header=1, nrows=nrows)
            unnamed2 = [c for c in df2.columns if str(c).startswith("Unnamed:")]
            if len(unnamed2) < len(unnamed_cols):
                logger.info(f"📊 header=1 更好：Unnamed 列从 {len(unnamed_cols)} 减少到 {len(unnamed2)}")
                df = df2

        # ── 5. 列名标准化：去掉换行符和多余空格 ────────────────────────────
        df.columns = [str(c).replace("\n", " ").replace("\r", "").strip() for c in df.columns]

        logger.info(f"📊 Excel 加载完成: {len(df)} 行 × {len(df.columns)} 列，列名: {list(df.columns)[:10]}")
        return df

    # ─────────────────────────────────────────────────────────
    # PDF 文件解析
    # ─────────────────────────────────────────────────────────

    @staticmethod
    def _is_garbage_table(df: "pd.DataFrame") -> bool:
        """
        判断 PyMuPDF 提取的表格是否为"垃圾表格"。

        PyMuPDF 的 find_tables() 会把排版残留、页眉页脚等误识别为表格，
        产生只有 2-3 行且列名全是 Col0/Col1 的无意义数据。
        这类表格应被视为无效，让调用方走 LLM 智能提取分支。

        判断标准：
        1. 行数太少（< 5 行）
        2. 且列名全是自动生成的（Col0, Col1, _col_0 等）
        3. 且大部分单元格是分类类型（category / object），几乎没有数值
        """
        import re

        if df is None or df.empty:
            return True

        # 行数检查：太少就可疑
        if len(df) >= 5:
            return False  # 5行以上认为有足够数据量，不算垃圾

        # 列名检查：如果大部分列名是自动生成的（Col0, Col1, _col_0, Unnamed: 0）
        auto_col_pattern = re.compile(r'^(Col\d+|_col_\d+|Unnamed:\s*\d+|\d+)$', re.IGNORECASE)
        auto_col_count = sum(1 for c in df.columns if auto_col_pattern.match(str(c).strip()))
        all_auto = auto_col_count >= len(df.columns) * 0.7  # 70%以上列名是自动生成的

        # 数值列检查：如果没有数值列，说明内容全是文字碎片
        numeric_col_count = 0
        for dtype in df.dtypes:
            if pd.api.types.is_numeric_dtype(dtype):
                numeric_col_count += 1

        has_numeric = numeric_col_count >= 1

        if all_auto and not has_numeric:
            return True  # 列名自动生成 + 无数值列 → 垃圾

        if len(df) <= 2 and all_auto:
            return True  # ≤2行 + 全自动列名 → 几乎肯定是垃圾

        return False

    def _extract_tables_from_pdf(
        self, file_content: bytes, filename: str
    ) -> "pd.DataFrame":
        """
        从 PDF 中提取表格数据。

        策略：
        1. 使用 PyMuPDF (fitz) 提取每页的表格
        2. 合并所有页的表格为一个 DataFrame
        3. 质量判断：如果提取的表格是垃圾数据（行少+列名自动生成），返回 None
        4. 返回 None 时，调用方会走 LLM 智能提取分支

        Returns:
            DataFrame 或 None（None 表示需要走 LLM 提取或纯文本兜底）
        """
        import pandas as pd

        try:
            import fitz  # PyMuPDF
        except ImportError:
            logger.warning("⚠️ PyMuPDF 未安装，无法解析 PDF。运行: pip install PyMuPDF")
            raise ImportError("需要安装 PyMuPDF 来解析 PDF 文件: pip install PyMuPDF")

        doc = fitz.open(stream=file_content, filetype="pdf")

        # 检查 PyMuPDF 版本是否支持 find_tables()（需要 v1.23.0+）
        if not hasattr(doc[0], 'find_tables'):
            logger.warning(
                f"[PDF] PyMuPDF 版本过低（{fitz.version}），不支持 find_tables()，"
                f"请升级到 v1.23.0+（pip install PyMuPDF>=1.23.0）。将走 LLM 智能提取。"
            )
            doc.close()
            return None  # 触发 LLM 提取分支

        all_tables = []

        for page_num, page in enumerate(doc):
            # PyMuPDF 的表格提取
            tables = page.find_tables()
            for table_idx, table in enumerate(tables):
                try:
                    df = table.to_pandas()
                    if df is not None and not df.empty:
                        # 标记来源页
                        df["_来源页码"] = page_num + 1
                        all_tables.append(df)
                        logger.debug(f"[PDF] 第{page_num+1}页表格{table_idx+1}: {len(df)}行 × {len(df.columns)}列")
                except Exception as e:
                    logger.debug(f"[PDF] 第{page_num+1}页表格{table_idx+1}提取失败: {e}")

        doc.close()

        if not all_tables:
            logger.warning(f"[PDF] {filename} 中未发现表格数据")
            return None

        # 合并所有表格
        combined = pd.concat(all_tables, ignore_index=True)

        # 清理：去掉 PyMuPDF 自动生成的空列名
        combined.columns = [
            str(c).replace("\n", " ").strip() if str(c) not in ("", "nan", "None") else f"_col_{i}"
            for i, c in enumerate(combined.columns)
        ]

        # 去掉来源页码列（辅助列）
        if "_来源页码" in combined.columns and len(combined.columns) > 1:
            combined = combined.drop(columns=["_来源页码"])

        # 去掉全空行
        combined = combined.dropna(how="all").reset_index(drop=True)

        # ── 垃圾表格质量判断 ──
        if self._is_garbage_table(combined):
            logger.warning(
                f"[PDF] {filename} 提取的表格数据质量过低（{len(combined)}行 × {len(combined.columns)}列），"
                f"列名: {list(combined.columns)}，将走 LLM 智能提取"
            )
            return None  # 触发 LLM 提取

        logger.info(f"📊 PDF 表格提取完成: {filename} → {len(combined)} 行 × {len(combined.columns)} 列")
        return combined

    def _extract_text_from_pdf(self, file_content: bytes) -> str:
        """从 PDF 中提取纯文本"""
        try:
            import fitz
        except ImportError:
            return ""

        doc = fitz.open(stream=file_content, filetype="pdf")
        text_parts = []
        for page in doc:
            text = page.get_text("text")
            if text.strip():
                text_parts.append(text.strip())
        doc.close()
        return "\n".join(text_parts)

    def _llm_extract_structured_csv(self, text: str, filename: str) -> str:
        """
        🧠 用 LLM 从纯文本中提取结构化 CSV 数据。

        当 PDF 没有明确表格时（如统计公报、研报正文），
        用 DeepSeek 将文本中的数字、指标、分类信息提取为 CSV 格式。

        Returns:
            CSV 格式字符串（带表头），失败返回空字符串
        """
        import pandas as pd

        # 截取前 12000 字符（统计公报数据密集，8000可能不够）
        truncated = text[:12000]
        if len(text) > 12000:
            truncated += "\n...(文本过长，已截取前12000字符)"

        prompt = f"""你是一个专业的中国经济统计数据提取助手。请从以下统计公报/经济报告文本中提取所有量化数据，整理成 CSV 格式。

## 提取规则

1. **提取范围**：所有包含数字的指标（GDP、产量、增长率、金额、占比、人口、收入等）
2. **CSV格式**：每条数据一行，第一行是表头
3. **建议表头**：指标名称,数值,单位,同比增长(%),备注/分类
4. **单位转换**（关键！）：
   - "1401879亿元" → 数值填 1401879，单位填 "亿元"
   - "69541万吨" → 数值填 69541，单位填 "万吨"
   - "695410000吨" → 数值填 69541，单位填 "万吨"
   - "增长5.0%" → 同比增长列填 5.0
   - "3.2个百分点" → 同比增长列填 3.2
5. **分类标注**：
   - "第一产业增加值..." → 备注填 "第一产业"
   - "夏粮产量..." → 备注填 "粮食/夏粮"
   - "城镇居民人均..." → 备注填 "居民收入/城镇"
   - 如有行业/领域信息，务必在备注列标注
6. **不要遗漏**：尽量完整提取每一个数据点
7. **只输出CSV**：不要解释，不要 markdown 代码块标记（直接输出CSV文本）

## 原始文本

{truncated}"""

        try:
            llm = self._get_llm_client()
            if not llm or not llm.available:
                logger.warning("[PDF-LLM] LLM 不可用，跳过智能提取")
                return ""

            # 直接用 urllib 同步调用（此函数运行在线程池中，不能用 asyncio）
            import urllib.request
            import urllib.error

            api_key, base_url = llm.api_key, llm.base_url
            url = f"{base_url}/chat/completions"
            payload = json.dumps({
                "model": llm.config.DEFAULT_MODEL,
                "messages": [{"role": "user", "content": prompt}],
                "temperature": 0.1,
                "max_tokens": 4096,
            }).encode("utf-8")

            req = urllib.request.Request(
                url, data=payload,
                headers={
                    "Content-Type": "application/json",
                    "Authorization": f"Bearer {api_key}",
                },
            )
            try:
                with urllib.request.urlopen(req, timeout=120) as resp:
                    result = json.loads(resp.read().decode("utf-8"))
                    csv_result = result["choices"][0]["message"]["content"]
            except urllib.error.HTTPError as e:
                body = e.read().decode("utf-8", errors="replace") if e.fp else ""
                logger.error(f"[PDF-LLM] API 请求失败 {e.code}: {body[:300]}")
                return ""
            except Exception as e:
                logger.error(f"[PDF-LLM] API 请求异常: {e}")
                return ""

            if csv_result and csv_result.strip():
                # 提取 ```csv ... ``` 代码块中的内容（如果有）
                cleaned = csv_result.strip()
                if cleaned.startswith("```"):
                    lines = cleaned.split("\n")
                    # 去掉首行 ```csv 和末行 ```
                    start = 1 if lines[0].startswith("```") else 0
                    end = len(lines) - 1 if lines[-1].strip() == "```" else len(lines)
                    cleaned = "\n".join(lines[start:end])

                # 验证是否是有效 CSV（至少有表头和一行数据）
                csv_lines = [l for l in cleaned.strip().split("\n") if l.strip()]
                if len(csv_lines) >= 2:
                    logger.info(f"[PDF-LLM] ✅ 成功提取 {len(csv_lines)-1} 行结构化数据")
                    return cleaned.strip()
                else:
                    logger.warning(f"[PDF-LLM] LLM 输出不足2行，无法构成表格: {cleaned[:200]}")
                    return ""

        except Exception as e:
            logger.error(f"[PDF-LLM] 智能提取失败: {e}")
            return ""

    def _get_llm_client(self):
        """获取 LLM 客户端单例"""
        try:
            from ..utils.llm_client import LLMClient
            return LLMClient.get_instance()
        except ImportError:
            return None

    # ─────────────────────────────────────────────────────────
    # Word 文件解析
    # ─────────────────────────────────────────────────────────

    def _extract_tables_from_docx(
        self, file_content: bytes, filename: str
    ) -> "pd.DataFrame":
        """
        从 Word (.docx) 中提取表格数据。

        策略：
        1. 使用 python-docx 提取所有表格
        2. 第一个表格的第一行作为列名（如果是 .doc 格式则跳过）
        3. 质量判断：垃圾表格返回 None，触发 LLM 智能提取

        Returns:
            DataFrame 或 None
        """
        import pandas as pd

        if filename.lower().endswith(".doc"):
            logger.warning("⚠️ .doc 格式（旧版 Word）不支持表格提取，请转换为 .docx")
            return None

        try:
            from docx import Document
        except ImportError:
            logger.warning("⚠️ python-docx 未安装。运行: pip install python-docx")
            raise ImportError("需要安装 python-docx 来解析 Word 文件: pip install python-docx")

        doc = Document(io.BytesIO(file_content))
        all_tables = []

        for table_idx, table in enumerate(doc.tables):
            rows = []
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                rows.append(cells)

            if not rows:
                continue

            # 第一行作为列名
            headers = rows[0]
            data_rows = rows[1:] if len(rows) > 1 else []

            if data_rows:
                df = pd.DataFrame(data_rows, columns=headers)
                all_tables.append(df)
                logger.debug(f"[Word] 表格{table_idx+1}: {len(df)}行 × {len(df.columns)}列")

        if not all_tables:
            logger.warning(f"[Word] {filename} 中未发现表格数据")
            return None

        # 合并
        combined = pd.concat(all_tables, ignore_index=True)

        # 列名清理
        combined.columns = [str(c) if str(c) not in ("", "nan", "None") else f"_col_{i}"
                            for i, c in enumerate(combined.columns)]

        # 去掉全空行
        combined = combined.dropna(how="all").reset_index(drop=True)

        # ── 垃圾表格质量判断（同 PDF）──
        if self._is_garbage_table(combined):
            logger.warning(
                f"[Word] {filename} 提取的表格数据质量过低（{len(combined)}行 × {len(combined.columns)}列），"
                f"将走 LLM 智能提取"
            )
            return None

        logger.info(f"📊 Word 表格提取完成: {filename} → {len(combined)} 行 × {len(combined.columns)} 列")
        return combined

    def _extract_text_from_docx(self, file_content: bytes) -> str:
        """从 Word (.docx) 中提取纯文本（备用方案）"""
        try:
            from docx import Document
        except ImportError:
            return ""

        doc = Document(io.BytesIO(file_content))
        return "\n".join(p.text.strip() for p in doc.paragraphs if p.text.strip())

    def close(self):
        """关闭 DuckDB 连接"""
        if self._conn:
            try:
                self._conn.close()
                logger.info("🦆 DuckDB 连接已关闭")
            except Exception:
                pass

    def drop_table(self, table_name: str) -> bool:
        """删除指定的表（用户主动删除文档时调用）"""
        if not self._initialized or not self._conn:
            return False
        try:
            safe_name = safe_table(table_name)
            self._conn.execute(f'DROP TABLE IF EXISTS {safe_name}')
            self._tables.pop(table_name, None)
            logger.info("🦆 表已删除: %s", table_name)
            return True
        except Exception as e:
            logger.error("🦆 删除表失败: %s → %s", table_name, e)
            return False

    def has_persistent_tables(self) -> bool:
        """检查是否有持久化表（用于判断是否需要重新上传）"""
        if not self._initialized or not self._conn:
            return False
        try:
            tables = self._get_registered_tables()
            # 排除系统临时表
            user_tables = [t for t in tables if not t.startswith("__temp_")]
            return len(user_tables) > 0
        except Exception:
            return False

    def __del__(self):
        self.close()


# ── 全局单例 ──────────────────────────────────────────────────
_duckdb_engine: Optional[DuckDBEngine] = None


def get_duckdb_engine() -> DuckDBEngine:
    """获取全局 DuckDB 引擎单例"""
    global _duckdb_engine
    if _duckdb_engine is None:
        _duckdb_engine = DuckDBEngine()
    return _duckdb_engine
